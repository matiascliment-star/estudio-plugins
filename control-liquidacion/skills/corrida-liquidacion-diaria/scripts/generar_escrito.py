#!/usr/bin/env python3
"""
Genera un DOCX de escrito judicial para la corrida diaria de caducidad.

Uso:
  python3 generar_escrito.py \
    --modelo pronto-despacho \
    --caratula "VALLEJOS, RAMIRO ALBERTO c/ SWISS MEDICAL ART S.A. s/DESPIDO" \
    --numero "CNT 000273/2024" \
    --placeholders '{"fecha_ultimo_hito":"07/07/2025","descripcion_ultimo_hito":"el Juzgado intimó al perito por remiso","estado_procesal":"paralizado en etapa probatoria","accion_especifica_a_pedir":"haga efectivo el apercibimiento al perito contador"}' \
    --output /tmp/pronto_VALLEJOS.docx

Aplica el FORMATO CANÓNICO importando desde
escritos-judiciales/scripts/formato_escrito.py.

NO duplica formato — los parámetros (fuente, sangrías, márgenes) viven todos
en el módulo canónico para que cambien en un solo lugar.
"""

import argparse
import json
import os
import re
import sys
from pathlib import Path

# Importar el helper canónico
HELPER_DIR = os.path.expanduser(
    '~/.claude/plugins/marketplaces/estudio-plugins/escritos-judiciales/scripts'
)
sys.path.insert(0, HELPER_DIR)

from formato_escrito import (  # noqa: E402
    nuevo_documento,
    titulo_principal,
    encabezado_tribunal,
    FUENTE,
    TAMANO_PT,
    SANGRIA_CUERPO_CM,
    SANGRIA_LETRADO_CM,
    LINEA_BLANCO_ANTES_SECCION_PT,
)
from docx.shared import Pt, Cm  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402


def _set_style(run, negrita=False, subrayado=False):
    run.font.name = FUENTE
    run.font.size = Pt(TAMANO_PT)
    run.bold = negrita
    run.underline = subrayado


def _parrafo_con_negritas_parciales(doc, texto, *, sangria_cm=SANGRIA_CUERPO_CM,
                                    bloques_negrita=None, space_before=0,
                                    negrita=False):
    """
    Párrafo justificado, fuente canónica, con tramos opcionales en negrita.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(sangria_cm) if sangria_cm else Cm(0)
    if space_before:
        pf.space_before = Pt(space_before)

    if bloques_negrita:
        pattern = "|".join(re.escape(b) for b in bloques_negrita)
        partes = re.split(f"({pattern})", texto)
        for parte in partes:
            if not parte:
                continue
            es_negrita = parte in bloques_negrita or negrita
            _set_style(p.add_run(parte), negrita=es_negrita)
    else:
        _set_style(p.add_run(texto), negrita=negrita)
    return p


def parsear_modelo(texto_modelo):
    """Devuelve la lista de bloques (separados por línea en blanco)."""
    lineas = [l.rstrip() for l in texto_modelo.strip().split("\n")]
    bloques, actual = [], []
    for l in lineas:
        if l.strip() == "":
            if actual:
                bloques.append("\n".join(actual).strip())
                actual = []
        else:
            actual.append(l)
    if actual:
        bloques.append("\n".join(actual).strip())
    return bloques


def reemplazar_placeholders(texto, placeholders):
    for k, v in placeholders.items():
        texto = texto.replace("{{" + k + "}}", str(v))
    sin_cubrir = re.findall(r"\{\{(\w+)\}\}", texto)
    if sin_cubrir:
        print(f"⚠️  Placeholders sin valor: {sin_cubrir}", file=sys.stderr)
    return texto


def generar(modelo_path, caratula, numero, placeholders, output_path):
    texto_modelo = Path(modelo_path).read_text(encoding="utf-8")
    placeholders = {**placeholders, "caratula": caratula, "numero": numero}
    texto = reemplazar_placeholders(texto_modelo, placeholders)
    bloques = parsear_modelo(texto)

    doc = nuevo_documento()

    for i, bloque in enumerate(bloques):
        # Bloque 0 — TÍTULO PRINCIPAL (justificado, negrita+subrayado, sin sangría)
        if i == 0:
            titulo_principal(doc, bloque)
            continue

        # Encabezado al tribunal corto ("Sr. Juez:", "Excmo. Tribunal:")
        if bloque.endswith(":") and len(bloque) < 50 and "\n" not in bloque:
            encabezado_tribunal(doc, bloque)
            continue

        # Cierre "SERÁ JUSTICIA"
        if "JUSTICIA" in bloque.upper() and len(bloque) < 40:
            _parrafo_con_negritas_parciales(
                doc, bloque, negrita=True,
                space_before=LINEA_BLANCO_ANTES_SECCION_PT,
            )
            continue

        # "Proveer de conformidad,"
        if bloque.lower().startswith("proveer") or bloque.lower().startswith("provea"):
            _parrafo_con_negritas_parciales(
                doc, bloque, space_before=LINEA_BLANCO_ANTES_SECCION_PT,
            )
            continue

        # Primer párrafo con datos del letrado (sangría 1.5 cm)
        if "GARCÍA CLIMENT" in bloque and "T°" in bloque:
            bloques_negrita = re.findall(r"\*\*(.+?)\*\*", bloque, flags=re.DOTALL)
            texto_limpio = re.sub(r"\*\*(.+?)\*\*", r"\1", bloque, flags=re.DOTALL)
            _parrafo_con_negritas_parciales(
                doc, texto_limpio,
                sangria_cm=SANGRIA_LETRADO_CM,
                bloques_negrita=bloques_negrita,
            )
            continue

        # Cuerpo normal (sangría 1.25 cm)
        bloques_negrita = re.findall(r"\*\*(.+?)\*\*", bloque, flags=re.DOTALL)
        texto_limpio = re.sub(r"\*\*(.+?)\*\*", r"\1", bloque, flags=re.DOTALL)
        _parrafo_con_negritas_parciales(
            doc, texto_limpio,
            sangria_cm=SANGRIA_CUERPO_CM,
            bloques_negrita=bloques_negrita,
        )

    doc.save(output_path)
    print(f"✅ Guardado: {output_path}")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--modelo", required=True, help="Nombre del modelo (sin extensión), ej 'pronto-despacho'")
    ap.add_argument("--caratula", required=True)
    ap.add_argument("--numero", required=True)
    ap.add_argument("--placeholders", required=True, help="JSON con placeholders del modelo")
    ap.add_argument("--output", required=True)
    args = ap.parse_args()

    skill_dir = Path(__file__).resolve().parent.parent
    modelo_path = skill_dir / "modelos" / f"{args.modelo}.md"
    if not modelo_path.exists():
        print(f"❌ Modelo no encontrado: {modelo_path}", file=sys.stderr)
        sys.exit(1)

    placeholders = json.loads(args.placeholders)
    generar(str(modelo_path), args.caratula, args.numero, placeholders, args.output)


if __name__ == "__main__":
    main()
