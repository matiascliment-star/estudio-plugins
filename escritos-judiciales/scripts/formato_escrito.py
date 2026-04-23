"""
Formato canónico para escritos judiciales DOCX del Estudio García Climent.

FUENTE DE VERDAD ÚNICA. Todos los skills del repo estudio-plugins que generan
DOCX deben importar desde acá. NO duplicar formato en cada skill.

Spec original: ~/.claude/projects/-Users-matiaschristiangarciacliment/memory/feedback_formato_escritos.md

Uso típico:

    from formato_escrito import (
        nuevo_documento,
        titulo_principal,
        encabezado_tribunal,
        parrafo_letrado,
        titulo_seccion,
        parrafo,
        firma,
    )

    doc = nuevo_documento()
    titulo_principal(doc, "INTERPONE RECURSO EXTRAORDINARIO FEDERAL")
    encabezado_tribunal(doc, "Excma. Corte Suprema de Justicia de la Nación:")
    parrafo_letrado(
        doc,
        "MATÍAS CHRISTIAN GARCÍA CLIMENT",
        "abogado, T° 97 F° 16 C.P.A.C.F., en autos ",
        "VÁZQUEZ, MIGUEL ANGEL c/ SWISS MEDICAL ART s/ ACCIDENTE - LEY ESPECIAL Expte. N° 045419/2021",
        ", a V.E. respetuosamente digo:",
    )
    titulo_seccion(doc, "I. OBJETO")
    parrafo(doc, "Vengo por el presente a interponer...")
    titulo_seccion(doc, "II. ANTECEDENTES")
    ...
    firma(doc)
    doc.save("/path/al/escrito.docx")

Para invocar como script (sanity check):
    python3 formato_escrito.py /tmp/sample.docx
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------------------------------------------------------------------------
# Configuración global
# ---------------------------------------------------------------------------

FUENTE = "Times New Roman"
TAMANO_PT = 12
INTERLINEADO = 1.5

MARGEN_SUP_CM = 2.0
MARGEN_INF_CM = 2.0
MARGEN_IZQ_CM = 3.0
MARGEN_DER_CM = 2.0

SANGRIA_CUERPO_CM = 1.25
SANGRIA_LETRADO_CM = 1.5

LINEA_BLANCO_ANTES_SECCION_PT = 12  # = 1 línea en blanco visual


# ---------------------------------------------------------------------------
# Builders
# ---------------------------------------------------------------------------

def nuevo_documento():
    """Documento con márgenes, fuente, tamaño e interlineado base."""
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(MARGEN_SUP_CM)
        s.bottom_margin = Cm(MARGEN_INF_CM)
        s.left_margin = Cm(MARGEN_IZQ_CM)
        s.right_margin = Cm(MARGEN_DER_CM)

    style = doc.styles["Normal"]
    style.font.name = FUENTE
    style.font.size = Pt(TAMANO_PT)
    style.paragraph_format.line_spacing = INTERLINEADO
    style.paragraph_format.space_after = Pt(0)
    return doc


def titulo_principal(doc, texto):
    """
    Título del escrito (objeto). JUSTIFICADO, sin sangría, negrita + subrayado.
    Ej: "INTERPONE RECURSO EXTRAORDINARIO FEDERAL".
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(texto.upper())
    run.bold = True
    run.underline = True
    run.font.name = FUENTE
    run.font.size = Pt(TAMANO_PT)
    return p


def encabezado_tribunal(doc, texto):
    """
    Encabezado de invocación al tribunal.
    Alineado a la IZQUIERDA, sin sangría.
    Ej: "Sr. Juez:", "Excma. Corte Suprema de Justicia de la Nación:".
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(texto)
    run.font.name = FUENTE
    run.font.size = Pt(TAMANO_PT)
    return p


def parrafo_letrado(doc, nombre, texto_pre_caratula, caratula, texto_post_caratula):
    """
    Primer párrafo del cuerpo (presentación del letrado).
    Sangría primera línea ~1.5 cm. Nombre del letrado y carátula van en NEGRITA.

    Ejemplo:
        parrafo_letrado(
            doc,
            "MATÍAS CHRISTIAN GARCÍA CLIMENT",
            ", abogado, T° 97 F° 16 C.P.A.C.F., con domicilio electrónico ..., ",
            "en autos VÁZQUEZ, MIGUEL ANGEL c/ ... Expte. N° 045419/2021",
            ", a V.S. respetuosamente digo:",
        )
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(SANGRIA_LETRADO_CM)

    r1 = p.add_run(nombre)
    r1.bold = True
    r1.font.name = FUENTE
    r1.font.size = Pt(TAMANO_PT)

    r2 = p.add_run(texto_pre_caratula)
    r2.font.name = FUENTE
    r2.font.size = Pt(TAMANO_PT)

    r3 = p.add_run(caratula)
    r3.bold = True
    r3.font.name = FUENTE
    r3.font.size = Pt(TAMANO_PT)

    r4 = p.add_run(texto_post_caratula)
    r4.font.name = FUENTE
    r4.font.size = Pt(TAMANO_PT)
    return p


def titulo_seccion(doc, texto):
    """
    Título de sección (I. OBJETO, II. HECHOS, etc.).
    Justificado, MISMA SANGRÍA que el cuerpo (1.25 cm), negrita + subrayado.
    Deja una línea en blanco antes (space_before = 12pt).
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(SANGRIA_CUERPO_CM)
    p.paragraph_format.space_before = Pt(LINEA_BLANCO_ANTES_SECCION_PT)
    run = p.add_run(texto)
    run.bold = True
    run.underline = True
    run.font.name = FUENTE
    run.font.size = Pt(TAMANO_PT)
    return p


def parrafo(doc, texto, sangria=True):
    """
    Párrafo de cuerpo. Justificado, sangría 1.25 cm en primera línea.
    Pasar sangria=False para bullets/petitorio sin sangría.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(SANGRIA_CUERPO_CM if sangria else 0)
    run = p.add_run(texto)
    run.font.name = FUENTE
    run.font.size = Pt(TAMANO_PT)
    return p


def firma(doc, nombre="MATÍAS CHRISTIAN GARCÍA CLIMENT",
         linea2="ABOGADO",
         linea3="T° 97 F° 16 C.P.A.C.F. / T° 46 F° 393 C.A.S.I."):
    """
    Bloque de firma. Centrado. Nombre en negrita.
    """
    doc.add_paragraph()  # línea en blanco antes
    for i, linea in enumerate((nombre, linea2, linea3)):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(linea)
        if i == 0:
            run.bold = True
        run.font.name = FUENTE
        run.font.size = Pt(TAMANO_PT)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run("(Firmado electrónicamente)")
    run.italic = True
    run.font.name = FUENTE
    run.font.size = Pt(TAMANO_PT)


# ---------------------------------------------------------------------------
# Sanity check CLI
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import sys
    out = sys.argv[1] if len(sys.argv) > 1 else "/tmp/sample_formato.docx"

    doc = nuevo_documento()
    titulo_principal(doc, "INTERPONE RECURSO EXTRAORDINARIO FEDERAL")
    encabezado_tribunal(doc, "Excma. Corte Suprema de Justicia de la Nación:")
    parrafo_letrado(
        doc,
        "MATÍAS CHRISTIAN GARCÍA CLIMENT",
        ", abogado, T° 97 F° 16 C.P.A.C.F., con domicilio electrónico constituido en 20-XXXXXXXX-X, en autos ",
        "VÁZQUEZ, MIGUEL ANGEL c/ SWISS MEDICAL ART s/ ACCIDENTE - LEY ESPECIAL Expte. N° 045419/2021",
        ", a V.E. respetuosamente digo:",
    )
    titulo_seccion(doc, "I. OBJETO")
    parrafo(doc, "Vengo por el presente a interponer recurso extraordinario federal en los términos del art. 14 de la Ley 48 contra la sentencia dictada en autos.")
    titulo_seccion(doc, "II. ANTECEDENTES")
    parrafo(doc, "El siniestro ocurrió el día...")
    titulo_seccion(doc, "III. PETITORIO")
    parrafo(doc, "Por todo lo expuesto, solicito a V.E.:", sangria=False)
    parrafo(doc, "1. Tenga por interpuesto el recurso extraordinario federal.", sangria=False)
    parrafo(doc, "2. Conceda el recurso y eleve los autos a la CSJN.", sangria=False)
    firma(doc)
    doc.save(out)
    print(f"OK -> {out}")
