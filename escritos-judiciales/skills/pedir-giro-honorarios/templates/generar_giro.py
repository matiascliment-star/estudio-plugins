"""
Template: generador de escrito de pedido de giro de honorarios (PJN/CABA).

USO:
  1. Copiar este archivo a /tmp/ y editarlo con los datos del caso concreto.
  2. Completar el dict CASO (ver abajo).
  3. Ejecutar: python3 /tmp/generar_giro.py
  4. El DOCX se guarda en ~/Desktop/{CARATULA_SHORT}_giro_honorarios.docx

REQUISITOS:
  - Las imágenes embebidas se toman de:
      ~/.claude/plugins/marketplaces/estudio-plugins/escritos-judiciales/skills/pedir-giro-honorarios/assets/constancia_afip.png
      ~/.claude/plugins/marketplaces/estudio-plugins/escritos-judiciales/skills/pedir-giro-honorarios/assets/constancia_cbu.png

FORMATO:
  Importa todos los builders desde el helper canónico
  (escritos-judiciales/scripts/formato_escrito.py).
  NO duplica formato — eso garantiza consistencia con el resto de los skills.
"""

import os
import sys

# Importar el helper canónico
HELPER_DIR = os.path.expanduser(
    '~/.claude/plugins/marketplaces/estudio-plugins/escritos-judiciales/scripts'
)
sys.path.insert(0, HELPER_DIR)

from formato_escrito import (  # noqa: E402
    nuevo_documento,
    titulo_principal,
    encabezado_tribunal,
    titulo_seccion,
    parrafo,
    FUENTE,
    TAMANO_PT,
    SANGRIA_LETRADO_CM,
)
from docx.shared import Pt, Cm, Inches  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402

# ======================================================================
# DATOS DEL CASO — EDITAR
# ======================================================================
CASO = {
    # Datos del expediente
    'caratula': 'APELLIDO, NOMBRE C/ DEMANDADA S/RECURSO LEY 27348',
    'numero_expte': 'CNT 000000/2025',
    'caratula_short': 'apellido',  # para nombre de archivo

    # Variante: 'A' (giro simple) o 'B' (traba embargo + giro)
    'variante': 'A',

    # Montos
    'honorarios': '1.000.000,00',
    'iva': '210.000,00',
    'monto_total': '1.210.000,00',

    # Fechas de depósitos
    'fecha_deposito': '23/10/2022',
    'fecha_embargo': '12/03/2026',
    'fechas_depositos': '07/11/2025 y 11/04/2026',

    # Contexto para variante B
    'recurso_pendiente': 'queja ante la CSJN',

    # Datos fijos del suscripto
    'letrado_nombre': 'MATÍAS CHRISTIAN GARCÍA CLIMENT',
    'letrado_tomo_folio': 'T° 97 F° 16 del C.P.A.C.F.',
    'letrado_cuit': '20-31380619-8',
    'letrado_dni': '31.380.619',
    'letrado_domicilio': 'Av. Ricardo Balbín 2368, CABA',
    'letrado_zona': '204',
    'letrado_email': 'matiasgarciacliment@gmail.com',
    'letrado_tel': '4-545-2488',
    'letrado_de': '2031306198',
    'banco': 'Banco de la Ciudad de Buenos Aires',
    'caja_ahorro': '000000260200356738',
    'cbu': '0290026110000003567389',
}

ASSETS_DIR = os.path.expanduser(
    '~/.claude/plugins/marketplaces/estudio-plugins/escritos-judiciales/skills/pedir-giro-honorarios/assets'
)

# ======================================================================
# HELPERS LOCALES (sólo lo que el módulo canónico no cubre)
# ======================================================================

def parrafo_mixto(doc, segs, indent_cm=1.25):
    """
    Párrafo justificado con runs mixtos (negritas/subrayados parciales).
    Cada seg es (texto, bold) o (texto, bold, underline).
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(indent_cm)
    for seg in segs:
        if len(seg) == 3:
            t, b, u = seg
        else:
            t, b = seg; u = False
        r = p.add_run(t)
        r.font.name = FUENTE
        r.font.size = Pt(TAMANO_PT)
        r.bold = b
        r.underline = u
    return p


def parrafo_centrado(doc, texto, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(texto)
    r.font.name = FUENTE
    r.font.size = Pt(TAMANO_PT)
    r.bold = bold
    return p


def constancia(doc, titulo, imagen, width_inches):
    doc.add_page_break()
    parrafo_centrado(doc, titulo, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run().add_picture(imagen, width=Inches(width_inches))


# ======================================================================
# CONSTRUIR EL ESCRITO
# ======================================================================
doc = nuevo_documento()

# --- TÍTULO PRINCIPAL ----------------------------------------------------
if CASO['variante'] == 'B':
    titulo = ('TRABA EMBARGO SOBRE FONDOS DEPOSITADOS. SE LIBRE GIRO ELECTRÓNICO '
              'EN CONCEPTO DE HONORARIOS. ADJUNTA DOCUMENTACIÓN. DECLARA BAJO '
              'JURAMENTO. HACE RESERVA.-')
else:
    titulo = ('SE LIBRE GIRO ELECTRÓNICO EN CONCEPTO DE HONORARIOS. ADJUNTA '
              'DOCUMENTACIÓN. DECLARA BAJO JURAMENTO. HACE RESERVA.-')
titulo_principal(doc, titulo)
doc.add_paragraph()
doc.add_paragraph()

# --- ENCABEZADO ---------------------------------------------------------
encabezado_tribunal(doc, "Sr. Juez:")

# --- PRIMER PÁRRAFO (LETRADO) -------------------------------------------
parrafo_mixto(doc, [
    (CASO['letrado_nombre'], True),
    (f', abogado, inscripto en el {CASO["letrado_tomo_folio"]}, C.U.I.T N° {CASO["letrado_cuit"]}, '
     f'IVA responsable inscripto, ', False),
    ('POR DERECHO PROPIO', True),
    (f', manteniendo el domicilio procesal en la {CASO["letrado_domicilio"]} (zona de notificación '
     f'{CASO["letrado_zona"]}, e-mail: {CASO["letrado_email"]}, tel: {CASO["letrado_tel"]}) y '
     f'domicilio electrónico en {CASO["letrado_de"]}, en los autos caratulados ', False),
    (f'"{CASO["caratula"]}" Expte. N° {CASO["numero_expte"]}', True),
    (', a V.S. digo:', False),
], indent_cm=SANGRIA_LETRADO_CM)

# --- I. (depende de variante) -------------------------------------------
if CASO['variante'] == 'B':
    titulo_seccion(doc, 'I.- TRABA EMBARGO SOBRE FONDOS DEPOSITADOS. SE LIBRE GIRO ELECTRÓNICO.-')
    parrafo_mixto(doc, [
        ('VISTO', True),
        (f' los depósitos efectuados por la demandada en fechas {CASO["fechas_depositos"]}, '
         f'que totalizan la suma de ${CASO["monto_total"]} (honorarios ${CASO["honorarios"]} + '
         f'IVA ${CASO["iva"]}) en la cuenta de autos —los que se encuentran dados en embargo '
         f'por la propia demandada con pretensión de inmovilización en plazo fijo hasta la '
         f'resolución de la {CASO["recurso_pendiente"]}—, y atento a que dicho trámite no '
         'suspende la ejecución ni la exigibilidad de los honorarios firmes (art. 285 CPCCN), ',
         False),
        ('SOLICITO', True),
        (' a V.S. que:', False),
    ])
    parrafo(doc,
        f'(a) Se TRABE EMBARGO a favor del suscripto, {CASO["letrado_nombre"]}, CUIT '
        f'{CASO["letrado_cuit"]}, sobre las sumas depositadas en la cuenta de autos por la '
        f'parte demandada, hasta cubrir la totalidad de ${CASO["monto_total"]} en concepto de '
        'honorarios profesionales, con más lo que corresponda por diferencias pendientes '
        'conforme las reservas oportunamente formuladas.'
    )
    parrafo_mixto(doc, [
        ('(b) Se ', False),
        ('ORDENE TRANSFERIR', True),
        (' la suma de ', False),
        (f'${CASO["monto_total"]}', True),
        (' desde la cuenta de autos a la cuenta bancaria del suscripto, ', False),
        (CASO['letrado_nombre'], True),
        (f', DNI {CASO["letrado_dni"]}, CUIT {CASO["letrado_cuit"]}, abierta en el '
         f'{CASO["banco"]}, Caja de Ahorro $ N° {CASO["caja_ahorro"]}, CBU {CASO["cbu"]}, '
         'bajo estricta responsabilidad del profesional firmante.', False),
    ])
else:
    titulo_seccion(doc, 'I.- SE ORDENE TRANSFERENCIA.-')
    parrafo_mixto(doc, [
        ('VISTO', True),
        (f' el depósito efectuado por la demandada en fecha {CASO["fecha_deposito"]} y el '
         f'embargo ordenado sobre la cuenta de autos en fecha {CASO["fecha_embargo"]}, ',
         False),
        ('SOLICITO', True),
        (' a S.S. que ', False),
        ('ORDENE TRANSFERIR', True),
        (' la suma de ', False),
        (f'${CASO["monto_total"]}', True),
        (' (honorarios ', False),
        (f'${CASO["honorarios"]}', True),
        (' + IVA ', False),
        (f'${CASO["iva"]}', True),
        (') desde la cuenta de autos a la cuenta bancaria del suscripto, ', False),
        (CASO['letrado_nombre'], True),
        (f', DNI {CASO["letrado_dni"]}, CUIT {CASO["letrado_cuit"]}, abierta en el '
         f'{CASO["banco"]}, Caja de Ahorro $ {CASO["caja_ahorro"]}, CBU {CASO["cbu"]}.',
         False),
    ])

# --- II. ADJUNTA DOCUMENTACIÓN ------------------------------------------
titulo_seccion(doc, 'II.- ADJUNTA DOCUMENTACIÓN. DECLARA BAJO JURAMENTO.-')
parrafo(doc,
    'Adjunto constancia de inscripción y constancia de CBU de la cuenta bancaria del '
    'suscripto y declaro bajo juramento que ambas piezas son auténticas.'
)

# --- III. PRESTA JURAMENTO ----------------------------------------------
titulo_seccion(doc, 'III.- PRESTA JURAMENTO.-')
parrafo(doc, 'Presto juramento de haber sido el único letrado actuante en los presentes autos.')

# --- IV. HACE RESERVA ---------------------------------------------------
titulo_seccion(doc, 'IV.- HACE RESERVA.-')
parrafo(doc,
    'Esta parte se reserva el derecho a reclamar los intereses devengados por la falta de '
    'pago en término, dejando expresa constancia que las sumas recibidas se imputarán en '
    'primer lugar a intereses y en segundo lugar a honorarios. En efecto, en los términos '
    'del artículo 900 del CCCN se deja constancia que el suscripto no presta su consentimiento '
    'a que la suma percibida se impute, en primer término, a la deuda principal de honorarios.'
)
parrafo(doc,
    'Asimismo, esta parte se reserva el derecho a reclamar la diferencia de honorarios que '
    'surja de aplicar lo dispuesto en el artículo 51 de la Ley 27.423 que dice:'
)
parrafo_mixto(doc, [
    ('"La regulación de honorarios deberá contener, bajo pena de nulidad, el monto expresado '
     'en moneda de curso legal y la cantidad de UMA que éste representa a la fecha de la '
     'resolución. El pago será definitivo y cancelatorio únicamente si se abona la cantidad '
     'de moneda de curso legal que resulte equivalente a la cantidad de UMA contenidas en la '
     'resolución regulatoria, ', False),
    ('según su valor vigente al momento del pago', False, True),
    ('".', False),
])
parrafo(doc, 'Solicito se tenga presente lo manifestado.')

# --- CIERRE -------------------------------------------------------------
doc.add_paragraph()
parrafo_centrado(doc, 'Proveer de conformidad,')
parrafo_centrado(doc, 'SERÁ JUSTICIA.-', bold=True)

# --- ANEXOS -------------------------------------------------------------
constancia(doc, 'CONSTANCIA DE INSCRIPCIÓN AFIP/ARCA',
           os.path.join(ASSETS_DIR, 'constancia_afip.png'), 6)
constancia(doc, 'CONSTANCIA DE CBU - BANCO CIUDAD',
           os.path.join(ASSETS_DIR, 'constancia_cbu.png'), 5)

# --- GUARDAR ------------------------------------------------------------
out = os.path.expanduser(f'~/Desktop/{CASO["caratula_short"]}_giro_honorarios.docx')
doc.save(out)
print(f'Guardado: {out}')
