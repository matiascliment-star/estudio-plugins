"""
Template: actualizar liquidación y reclamar diferencias (PJN/CABA).

USO:
  1. Copiar a /tmp/generar_actualizacion_<caso>.py y editar el dict CASO.
  2. python3 /tmp/generar_actualizacion_<caso>.py
  3. Output: ~/Desktop/{caratula_short}_nueva_liquidacion.docx

SOPORTA:
  - Capital con actualización RIPTE (solo factor)
  - Honorarios al valor UMA actual, con opción de escenarios (con/sin prorrateo)
  - Discriminación honorarios + IVA
  - Descuento de pagos parciales efectivos
  - Dos escenarios (principal/subsidiario) si hay prorrateo pendiente

NO SOPORTA aún (agregar a mano si hace falta):
  - Tasa activa capitalizable (ver modelo Valdez — se hace cálculo externo con CPACF)
  - IPC + tasa pura
  - CER
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ======================================================================
# DATOS DEL CASO — EDITAR
# ======================================================================
CASO = {
    # --- Expediente ---
    'caratula': 'APELLIDO, NOMBRE C/ DEMANDADA S/RECURSO LEY 27348',
    'numero_expte': 'CNT 000000/2025',
    'caratula_short': 'apellido',

    # --- Datos del letrado (fijos) ---
    'letrado_nombre': 'MATÍAS CHRISTIAN GARCÍA CLIMENT',
    'letrado_tomo_folio': 'T° 97 F° 16 del CPACF',
    'letrado_cuit': '20-31380619-8',
    'letrado_domicilio': 'Av. Ricardo Balbín 2368, CABA',
    'letrado_zona': '204',
    'letrado_email': 'matiasgarciacliment@gmail.com',
    'letrado_tel': '4-545-2488',
    'letrado_de': '2031306198',

    # --- CAPITAL: incluir o no este bloque ---
    'incluir_capital': True,
    'credito_original': '$8.519.357,00',
    'fecha_hecho': '28/07/2021',
    'metodo_actualizacion': 'RIPTE',  # 'RIPTE' | 'TASA_ACTIVA' | 'IPC_TASA_PURA' | 'CER'
    'ripte_inicial_label': 'RIPTE 07/2021',
    'ripte_inicial_valor': '10.089,96',
    'ripte_final_label': 'RIPTE 12/2025',
    'ripte_final_valor': '186.718,83',
    'factor_actualizacion': 'x 18,5054',
    'capital_actualizado': '$157.654.177,46',
    'fecha_index_final': '12/2025',
    'percibido_inc': '-$31.440.421,20',
    'percibido_inc_label': 'Menos lo percibido en Incidente N° 2',
    'pago_capital_label': 'Menos lo percibido el 07/11/2025 (capital)',
    'pago_capital_valor': '-$116.249.094,47',
    'saldo_capital': '$9.964.661,79',
    'explicacion_capital': (
        'La diferencia de $9.964.661,79 corresponde a la actualización por índice RIPTE '
        'devengada desde 08/2025 (fecha de la liquidación aprobada) hasta 12/2025 (fecha en '
        'que el actor percibió el último giro). Dicho saldo se imputa al actor y queda '
        'sujeto a nueva actualización a la fecha del efectivo pago.'
    ),

    # --- HONORARIOS: incluir o no este bloque ---
    'incluir_honorarios': True,
    'hay_prorrateo': True,  # Si True, genera dos escenarios
    'uma_actual': '$92.482,00',
    'uma_actual_valor': 92482.00,
    'uma_resolucion': 'Res. 538/26 CSJN (vigente feb-2026)',

    # Honorarios 1ra instancia
    'uma_1ra': 375.00,
    'uma_1ra_prorr': 363.72,  # si hay prorrateo

    # Honorarios REX (no se prorratea)
    'uma_rex': 48.75,

    # Pago recibido
    'fecha_pago_hon': '11/04/2026',
    'pagado_total_con_iva': 27980223.98,
    # Si no se sabe discriminar, se calcula automáticamente /1.21

    # Nota sobre imputaciones tramposas de la demandada (opcional)
    'nota_imputacion_tramposa': (
        'Importante: a diferencia de lo pretendido por la demandada en su escrito del '
        '11/04/2026 —que descontó falsamente $6.086.865,54 como "dado en embargo el '
        '07/11/2025"—, el único depósito efectivamente efectuado en concepto de honorarios '
        'del suscripto es la transferencia del 11/04/2026 por $27.980.223,98 (con IVA '
        'incluido). El depósito del 07/11/2025 por $116.249.094,47 fue imputado íntegramente '
        'al CAPITAL de condena, conforme surge del comprobante bancario.'
    ),

    # --- Contexto introductorio (párrafo I) ---
    'intro_i': (
        'Habiendo transcurrido tiempo desde la liquidación aprobada el 05/11/2025 —que sólo '
        'cubría la actualización hasta el 08/2025— y habiéndose producido pagos parciales '
        'desde entonces, esta parte viene a practicar nueva liquidación del capital y de '
        'los honorarios del suscripto conforme los parámetros de la sentencia firme, a fin '
        'de reclamar las DIFERENCIAS devengadas que a la fecha permanecen impagas.'
    ),
}


# ======================================================================
# CONSTRUCTOR
# ======================================================================
doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2); s.bottom_margin = Cm(2)
    s.left_margin = Cm(3); s.right_margin = Cm(2)

st = doc.styles['Normal']
st.font.name = 'Times New Roman'; st.font.size = Pt(12)
pf = st.paragraph_format
pf.line_spacing = 1.5
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
pf.first_line_indent = Cm(1.25)


def _run(p, t, bold=False, underline=False):
    r = p.add_run(t)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)
    r.bold = bold; r.underline = underline
    return r


def add_titulo(t):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    _run(p, t, bold=True, underline=True)


def add_encabezado(t):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    _run(p, t, bold=True)


def add_seccion(t):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    _run(p, t, bold=True, underline=True)


def add_par(t, indent=1.25, bold=False, align='justify'):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = {'justify': WD_ALIGN_PARAGRAPH.JUSTIFY, 'center': WD_ALIGN_PARAGRAPH.CENTER, 'left': WD_ALIGN_PARAGRAPH.LEFT}[align]
    p.paragraph_format.first_line_indent = Cm(indent) if indent else Cm(0)
    p.paragraph_format.line_spacing = 1.5
    _run(p, t, bold=bold)


def add_par_mixto(segs, indent=1.25):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(indent)
    p.paragraph_format.line_spacing = 1.5
    for seg in segs:
        t, b = seg[:2]; u = seg[2] if len(seg) > 2 else False
        _run(p, t, bold=b, underline=u)


def set_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for b in ('top','left','bottom','right'):
        border = OxmlElement(f'w:{b}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_tabla(headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text=''
        p = c.paragraphs[0]
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.font.name='Times New Roman'; r.font.size=Pt(11); r.bold=True
        set_cell_borders(c)
    for ri, row in enumerate(rows, start=1):
        for ci, v in enumerate(row):
            c = t.rows[ri].cells[ci]; c.text=''
            p = c.paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci==0 else WD_ALIGN_PARAGRAPH.RIGHT
            txt = v; bold = False
            if isinstance(txt, str) and txt.startswith('**') and txt.endswith('**'):
                txt = txt[2:-2]; bold = True
            r = p.add_run(str(txt)); r.font.name='Times New Roman'; r.font.size=Pt(11); r.bold=bold
            set_cell_borders(c)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = w
    return t


def fmt(n):
    """Formato argentino: miles con punto, decimales con coma, dos decimales."""
    return f"${n:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')


# ======================================================================
# TÍTULO
# ======================================================================
tit = 'PARTE ACTORA PRACTICA NUEVA LIQUIDACIÓN'
if CASO['incluir_capital'] and CASO['incluir_honorarios']:
    tit += ' DE CAPITAL Y HONORARIOS'
elif CASO['incluir_honorarios']:
    tit += ' DE HONORARIOS'

met = []
if CASO['incluir_capital']:
    met.append('ÍNDICE ' + CASO['metodo_actualizacion'])
if CASO['incluir_honorarios']:
    met.append('VALOR UMA VIGENTE')
if met:
    tit += '. ACTUALIZA POR ' + ' Y '.join(met)
tit += '. RECLAMA DIFERENCIAS. SE CORRA TRASLADO.-'
add_titulo(tit)
doc.add_paragraph(); doc.add_paragraph()

# ======================================================================
# ENCABEZADO
# ======================================================================
add_encabezado('Sr. Juez:')

# ======================================================================
# PRIMER PÁRRAFO
# ======================================================================
add_par_mixto([
    ('', False, False),
    (CASO['letrado_nombre'], True, False),
    (f', abogado inscripto al {CASO["letrado_tomo_folio"]}, C.U.I.T {CASO["letrado_cuit"]}, '
     f'IVA responsable inscripto, apoderado de la PARTE ACTORA y POR DERECHO PROPIO, '
     f'manteniendo el domicilio procesal en la {CASO["letrado_domicilio"]} (zona de '
     f'notificación {CASO["letrado_zona"]}, e-mail: {CASO["letrado_email"]}, tel: '
     f'{CASO["letrado_tel"]}) y domicilio electrónico en {CASO["letrado_de"]}, en los autos '
     f'caratulados ', False, False),
    (f'"{CASO["caratula"]}" Expte. N° {CASO["numero_expte"]}', True, False),
    (', a V.S. digo:', False, False),
], indent=1.5)

# ======================================================================
# I - PRACTICA NUEVA LIQUIDACIÓN
# ======================================================================
add_seccion('I.- PRACTICA NUEVA LIQUIDACIÓN.-')
add_par(CASO['intro_i'])

# ---- BLOQUE A: CAPITAL ----
if CASO['incluir_capital']:
    add_par(f'A) CAPITAL DE CONDENA (actualización {CASO["metodo_actualizacion"]}):',
            indent=1.25, bold=True)
    add_tabla(
        headers=['Concepto', 'Valor'],
        rows=[
            [f'Crédito al {CASO["fecha_hecho"]}', CASO['credito_original']],
            [f'Índice {CASO["ripte_inicial_label"]}', CASO['ripte_inicial_valor']],
            [f'Índice {CASO["ripte_final_label"]}', CASO['ripte_final_valor']],
            ['Factor de actualización', CASO['factor_actualizacion']],
            [f'Capital actualizado al {CASO["fecha_index_final"]}', CASO['capital_actualizado']],
            [CASO['percibido_inc_label'], CASO['percibido_inc']],
            [CASO['pago_capital_label'], CASO['pago_capital_valor']],
            [f'**Saldo de CAPITAL pendiente al {CASO["fecha_index_final"]}**', f'**{CASO["saldo_capital"]}**'],
        ],
        widths=[Cm(10), Cm(5)]
    )
    doc.add_paragraph()
    add_par(CASO['explicacion_capital'])

# ---- BLOQUE B: HONORARIOS ----
if CASO['incluir_honorarios']:
    uma = CASO['uma_actual_valor']
    uma_str = CASO['uma_actual']
    pagado = CASO['pagado_total_con_iva']
    pagado_sin_iva = pagado / 1.21
    pagado_iva = pagado - pagado_sin_iva

    uma_1ra = CASO['uma_1ra']
    uma_rex = CASO['uma_rex']

    hon_1ra_sin = uma_1ra * uma
    hon_rex = uma_rex * uma
    sub_p = hon_1ra_sin + hon_rex
    iva_p = sub_p * 0.21
    tot_p = sub_p + iva_p
    saldo_p = tot_p - pagado

    add_par(f'B) HONORARIOS PROFESIONALES DEL SUSCRIPTO (valor UMA {CASO["fecha_index_final"]} / actual):',
            indent=1.25, bold=True)

    if CASO['hay_prorrateo']:
        uma_1ra_pr = CASO['uma_1ra_prorr']
        hon_1ra_con = uma_1ra_pr * uma
        sub_s = hon_1ra_con + hon_rex
        iva_s = sub_s * 0.21
        tot_s = sub_s + iva_s
        saldo_s = tot_s - pagado

        add_par(
            f'La liquidación de honorarios se efectúa en DOS ESCENARIOS: (1) principal, sin '
            f'prorrateo, con honorarios de 1ra instancia por {uma_1ra:g} UMAs; y (2) subsidiario, '
            f'con prorrateo, con honorarios de 1ra instancia prorrateados en {uma_1ra_pr} UMAs. '
            f'Los honorarios del REX ({uma_rex} UMAs) no se prorratean en ningún caso. Todas las '
            f'sumas se expresan al valor UMA vigente ({uma_str}, {CASO["uma_resolucion"]}).'
        )

    if CASO.get('nota_imputacion_tramposa'):
        add_par(CASO['nota_imputacion_tramposa'])
    doc.add_paragraph()

    # (1) Principal
    label_1 = '(1) Liquidación principal — SIN prorrateo:' if CASO['hay_prorrateo'] else 'Liquidación:'
    add_par(label_1, indent=1.25, bold=True)
    add_tabla(
        headers=['Concepto', 'UMAs', 'Pesos'],
        rows=[
            ['Honorarios 1ra instancia', f'{uma_1ra:g}', fmt(hon_1ra_sin)],
            ['Honorarios REX (10% s/ instancias ordinarias)', f'{uma_rex}', fmt(hon_rex)],
            ['Subtotal honorarios (sin IVA)', f'{uma_1ra + uma_rex:g}', fmt(sub_p)],
            ['IVA 21%', '—', fmt(iva_p)],
            ['Total honorarios + IVA', '—', fmt(tot_p)],
            [f'Menos lo pagado (transf. {CASO["fecha_pago_hon"]})', '—', '-' + fmt(pagado)],
            ['**Saldo adeudado HONORARIOS' + (' — Escenario principal' if CASO['hay_prorrateo'] else '') + '**',
             '—', f'**{fmt(saldo_p)}**'],
        ],
        widths=[Cm(8), Cm(2.5), Cm(4.5)]
    )
    doc.add_paragraph()

    # (2) Subsidiario si aplica
    if CASO['hay_prorrateo']:
        add_par('(2) Liquidación subsidiaria — CON prorrateo art. 8 ley 24.432:',
                indent=1.25, bold=True)
        add_tabla(
            headers=['Concepto', 'UMAs', 'Pesos'],
            rows=[
                ['Honorarios 1ra instancia (prorrateados)', f'{uma_1ra_pr}', fmt(hon_1ra_con)],
                ['Honorarios REX (no se prorratea)', f'{uma_rex}', fmt(hon_rex)],
                ['Subtotal honorarios (sin IVA)', f'{uma_1ra_pr + uma_rex:g}', fmt(sub_s)],
                ['IVA 21%', '—', fmt(iva_s)],
                ['Total honorarios + IVA', '—', fmt(tot_s)],
                [f'Menos lo pagado (transf. {CASO["fecha_pago_hon"]})', '—', '-' + fmt(pagado)],
                ['**Saldo adeudado HONORARIOS — Escenario subsidiario**', '—', f'**{fmt(saldo_s)}**'],
            ],
            widths=[Cm(8), Cm(2.5), Cm(4.5)]
        )
        doc.add_paragraph()

    # ---- RESUMEN FINAL ----
    if CASO['incluir_capital']:
        add_par('C) RESUMEN — TOTAL ADEUDADO:', indent=1.25, bold=True)
        cap_saldo_num = float(CASO['saldo_capital'].replace('$','').replace('.','').replace(',','.'))
        if CASO['hay_prorrateo']:
            add_tabla(
                headers=['Concepto', 'Escenario principal', 'Escenario subsidiario'],
                rows=[
                    ['Capital (diferencia actualización)', CASO['saldo_capital'], CASO['saldo_capital']],
                    [f'Honorarios al UMA actual {uma_str} (c/IVA)', fmt(saldo_p), fmt(saldo_s)],
                    ['**TOTAL ADEUDADO POR LA DEMANDADA**',
                     f'**{fmt(cap_saldo_num + saldo_p)}**',
                     f'**{fmt(cap_saldo_num + saldo_s)}**'],
                ],
                widths=[Cm(7), Cm(4), Cm(4)]
            )
        else:
            add_tabla(
                headers=['Concepto', 'Valor'],
                rows=[
                    ['Capital (diferencia actualización)', CASO['saldo_capital']],
                    [f'Honorarios al UMA actual (c/IVA)', fmt(saldo_p)],
                    ['**TOTAL ADEUDADO POR LA DEMANDADA**', f'**{fmt(cap_saldo_num + saldo_p)}**'],
                ],
                widths=[Cm(10), Cm(5)]
            )
        doc.add_paragraph()

add_par(
    'Se deja expresamente reservado el derecho a practicar nueva liquidación al momento del '
    'efectivo pago, con el índice y el valor UMA vigentes a esa fecha, así como a reclamar '
    'los intereses moratorios devengados por la falta de pago oportuno conforme art. 900 CCCN.'
)

# ======================================================================
# II - SE CORRA TRASLADO
# ======================================================================
add_seccion('II.- SE CORRA TRASLADO.-')
add_par(
    'Solicito se ordene el traslado del presente escrito por tres días a la parte demandada '
    'a fin de que deposite el monto resultante de ésta liquidación en la cuenta de estos '
    'autos o, en su caso, la impugne, practique la liquidación que estime corresponder y '
    'deposite el monto que arroje la misma en la cuenta indicada, todo ello bajo '
    'apercibimiento de que transcurrido el plazo sin que medie observación válida, quedará '
    'aprobada la liquidación efectuada por esta parte y se procederá a la ejecución.'
)
add_par(
    'Asimismo, solicito a V.S. que se le haga saber a la accionada que si su intención es '
    'detener el curso de las diferencias por actualización, deberá practicar nueva '
    'liquidación a la fecha del efectivo pago y acreditar el depósito de la cantidad '
    'resultante, tomando en consideración que se entenderá como fecha de efectivo pago '
    'aquella en la que la acreedora quede notificada automáticamente del auto que hace saber '
    'la dación en pago efectuada.'
)

# Cierre
doc.add_paragraph()
add_par('Proveer de conformidad,', indent=0, align='center')
add_par('SERÁ JUSTICIA.-', indent=0, align='center', bold=True)

out = os.path.expanduser(f'~/Desktop/{CASO["caratula_short"]}_nueva_liquidacion.docx')
doc.save(out)
print(f'Guardado: {out}')
